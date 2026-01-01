"""
PDF and Document Text Extractor for IIM Sambalpur GPT

Extracts text from PDFs and DOCX files in the 'pdf and docs' folder
and outputs them in the structured format used by the scraped_data pipeline.
"""

import os
import hashlib
import re
from pathlib import Path

# PDF extraction
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDF_SUPPORT = True
except ImportError:
    try:
        import PyPDF2
        PDF_SUPPORT = "pypdf2"
    except ImportError:
        PDF_SUPPORT = False

# DOCX extraction
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# XLSX extraction
try:
    import openpyxl
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False


# Configuration
INPUT_DIR = "pdf and docs"
OUTPUT_DIR = os.path.join("scraped_data", "txt")
INSTITUTION = "IIM Sambalpur"


def get_file_hash(filepath):
    """Generate MD5 hash of filename for consistent naming."""
    return hashlib.md5(filepath.encode('utf-8')).hexdigest()


def detect_document_type(filename, text):
    """Detect the type of document based on filename and content."""
    content = (filename + " " + (text[:2000] if text else "")).lower()
    
    if any(k in content for k in ['brochure', 'programme', 'program']):
        return 'brochure'
    elif any(k in content for k in ['course', 'outline', 'syllabus', 'curriculum']):
        return 'course_outline'
    elif any(k in content for k in ['manual', 'handbook', 'guide']):
        return 'manual'
    elif any(k in content for k in ['calendar', 'schedule', 'timetable']):
        return 'schedule'
    elif any(k in content for k in ['exam', 'test', 'midterm', 'assignment']):
        return 'exam_material'
    else:
        return 'document'


def extract_tags(text, filename):
    """Extract relevant tags from document."""
    if not text:
        return ""
    
    # Common academic keywords
    keywords = []
    content = (text + " " + filename).lower()
    
    tag_patterns = [
        ('Data Science', ['data science', 'dsai', 'ds&ai']),
        ('AI', ['artificial intelligence', ' ai ', 'machine learning']),
        ('MBA', ['mba', 'master of business']),
        ('Public Policy', ['public policy', 'governance']),
        ('Management', ['management', 'organizational']),
        ('Statistics', ['statistics', 'statistical']),
        ('Mathematics', ['mathematics', 'calculus', 'linear algebra']),
        ('Philosophy', ['philosophy', 'ethics']),
        ('Psychology', ['psychology', 'behavioral']),
        ('Executive', ['executive', 'working professionals']),
    ]
    
    for tag, patterns in tag_patterns:
        if any(p in content for p in patterns):
            keywords.append(tag)
    
    return ', '.join(keywords[:10])


def extract_pdf_pdfminer(filepath):
    """Extract text from PDF using pdfminer.six."""
    try:
        return pdf_extract_text(filepath)
    except Exception as e:
        print(f"  pdfminer error: {e}")
        return None


def extract_pdf_pypdf2(filepath):
    """Extract text from PDF using PyPDF2."""
    try:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return '\n\n'.join(text)
    except Exception as e:
        print(f"  PyPDF2 error: {e}")
        return None


def extract_docx(filepath):
    """Extract text from DOCX file."""
    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        return '\n\n'.join(paragraphs)
    except Exception as e:
        print(f"  DOCX error: {e}")
        return None


def extract_xlsx(filepath):
    """Extract text from XLSX file."""
    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
        all_text = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            all_text.append(f"=== Sheet: {sheet_name} ===\n")
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    all_text.append(' | '.join(row_values))
        
        return '\n'.join(all_text)
    except Exception as e:
        print(f"  XLSX error: {e}")
        return None


def save_extracted_text(filename, text, source_path):
    """Save extracted text in the structured format."""
    file_hash = get_file_hash(source_path)
    output_path = os.path.join(OUTPUT_DIR, f"doc_{file_hash}.txt")
    
    doc_type = detect_document_type(filename, text)
    tags = extract_tags(text, filename)
    
    # Clean up the text
    if text:
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("==================== SOURCE ====================\n")
        f.write(f"URL: file://{source_path}\n")
        f.write(f"PAGE_TITLE: {filename}\n")
        f.write(f"SOURCE_TYPE: {doc_type}\n")
        f.write(f"INSTITUTION: {INSTITUTION}\n")
        f.write(f"LAST_UPDATED: extracted_from_local_file\n\n")
        f.write("==================== TEXT ====================\n")
        f.write(text or "No text could be extracted")
        f.write("\n\n")
        f.write("==================== METADATA ====================\n")
        f.write(f"TAGS: {tags}\n")
        f.write(f"ORIGINAL_FILE: {filename}\n")
        f.write(f"CONFIDENCE_LEVEL: high\n")
        f.write("\n================================================\n\n")
    
    return output_path


def main():
    print("=" * 60)
    print("IIM Sambalpur Document Extractor")
    print("=" * 60)
    
    # Check dependencies
    print("\nDependency Check:")
    print(f"  PDF Support: {PDF_SUPPORT}")
    print(f"  DOCX Support: {DOCX_SUPPORT}")
    print(f"  XLSX Support: {XLSX_SUPPORT}")
    
    if not PDF_SUPPORT:
        print("\n⚠️  Install pdfminer.six or PyPDF2 for PDF support:")
        print("    pip install pdfminer.six python-docx openpyxl")
    
    # Ensure output directory exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # Get all files
    input_path = Path(INPUT_DIR)
    if not input_path.exists():
        print(f"\n❌ Input directory '{INPUT_DIR}' not found!")
        return
    
    files = list(input_path.iterdir())
    print(f"\nFound {len(files)} files to process.\n")
    
    success_count = 0
    fail_count = 0
    
    for filepath in files:
        if filepath.is_dir():
            continue
            
        filename = filepath.name
        ext = filepath.suffix.lower()
        
        print(f"Processing: {filename}")
        
        text = None
        
        if ext == '.pdf':
            if PDF_SUPPORT == True:
                text = extract_pdf_pdfminer(str(filepath))
            elif PDF_SUPPORT == "pypdf2":
                text = extract_pdf_pypdf2(str(filepath))
            else:
                print("  ⏭️  Skipped (no PDF library)")
                fail_count += 1
                continue
                
        elif ext == '.docx':
            if DOCX_SUPPORT:
                text = extract_docx(str(filepath))
            else:
                print("  ⏭️  Skipped (python-docx not installed)")
                fail_count += 1
                continue
                
        elif ext == '.xlsx':
            if XLSX_SUPPORT:
                text = extract_xlsx(str(filepath))
            else:
                print("  ⏭️  Skipped (openpyxl not installed)")
                fail_count += 1
                continue
        else:
            print(f"  ⏭️  Skipped (unsupported format: {ext})")
            continue
        
        if text and len(text.strip()) > 50:
            output_path = save_extracted_text(filename, text, str(filepath))
            char_count = len(text)
            print(f"  ✅ Extracted {char_count:,} characters -> {os.path.basename(output_path)}")
            success_count += 1
        else:
            print(f"  ⚠️  No meaningful text extracted")
            fail_count += 1
    
    print("\n" + "=" * 60)
    print(f"COMPLETE: {success_count} succeeded, {fail_count} failed/skipped")
    print("=" * 60)


if __name__ == "__main__":
    main()
